from fastapi import FastAPI, UploadFile, File
from minio import Minio
from docx import Document
import os
import uuid
import shutil
import re
from typing import List, Dict, Optional, Any
from dataclasses import dataclass, field
import json
from datetime import datetime

@dataclass
class ContentBlock:
    """内容块 - 表示文档中的一个结构化单元（题目、选项、文本等）"""
    id: str                           # 唯一标识，如 block_001
    type: str                         # 类型：question/options/text/table
    sequence: int                     # 在文档中的序号
    text: str                         # 文本内容
    image_refs: List[str] = field(default_factory=list)  # 关联的图片ID列表
    paragraph_index: int = 0          # 在原文档中的段落索引  
    style: Optional[str] = None       # 段落样式
    parent_block: Optional[str] = None # 父块ID（选项关联到题目）
    metadata: Dict[str, Any] = field(default_factory=dict)  # 额外元数据

@dataclass 
class ImageInfo:
    """图片信息"""
    id: str                          # 图片ID，如 img_001
    filename: str                    # 文件名
    url: str                         # MinIO URL
    type: str                        # 图片类型：docx_embedded/docx_shape等
    original_path: str               # 在DOCX中的原始路径
    paragraph_index: int = 0         # 所属段落索引
    position_in_paragraph: int = 0   # 在段落中的位置
    size: Optional[Dict[str, int]] = None  # 图片尺寸 {width, height}

@dataclass
class DoclingDocument:
    """结构化文档 - Docling格式的文档表示"""
    docling_version: str = "1.0"
    metadata: Dict[str, Any] = field(default_factory=dict)
    content_blocks: List[ContentBlock] = field(default_factory=list)
    images: Dict[str, ImageInfo] = field(default_factory=dict)
    relationships: Dict[str, List[str]] = field(default_factory=dict)  # 块与图片的关联关系
    
    def add_content_block(self, block: ContentBlock):
        """添加内容块"""
        self.content_blocks.append(block)
        
    def add_image(self, image: ImageInfo):
        """添加图片"""
        self.images[image.id] = image
        
    def link_image_to_block(self, block_id: str, image_id: str):
        """建立内容块与图片的关联"""
        # 找到对应的内容块并添加图片引用
        for block in self.content_blocks:
            if block.id == block_id:
                if image_id not in block.image_refs:
                    block.image_refs.append(image_id)
                break
        
        # 在关联关系中记录
        if block_id not in self.relationships:
            self.relationships[block_id] = []
        if image_id not in self.relationships[block_id]:
            self.relationships[block_id].append(image_id)
            
    def to_dict(self) -> Dict[str, Any]:
        """转换为字典格式"""
        return {
            "docling_version": self.docling_version,
            "metadata": self.metadata,
            "content_blocks": [
                {
                    "id": block.id,
                    "type": block.type,
                    "sequence": block.sequence,
                    "text": block.text,
                    "image_refs": block.image_refs,
                    "paragraph_index": block.paragraph_index,
                    "style": block.style,
                    "parent_block": block.parent_block,
                    "metadata": block.metadata
                }
                for block in self.content_blocks
            ],
            "images": {
                img_id: {
                    "id": img.id,
                    "filename": img.filename,
                    "url": img.url,
                    "type": img.type,
                    "original_path": img.original_path,
                    "paragraph_index": img.paragraph_index,
                    "position_in_paragraph": img.position_in_paragraph,
                    "size": img.size
                }
                for img_id, img in self.images.items()
            },
            "relationships": self.relationships
        }

app = FastAPI()

# MinIO configuration from environment variables
minio_endpoint = os.getenv("MINIO_ENDPOINT", "localhost:9000")
minio_access_key = os.getenv("MINIO_ACCESS_KEY")
minio_secret_key = os.getenv("MINIO_SECRET_KEY")
minio_secure = os.getenv("MINIO_SECURE", "false").lower() == "true"
minio_external_url = os.getenv("MINIO_EXTERNAL_URL", f"http://{minio_endpoint}")
bucket_name = os.getenv("MINIO_BUCKET", "exam-images")

minio_client = Minio(
    minio_endpoint,
    access_key=minio_access_key,
    secret_key=minio_secret_key,
    secure=minio_secure
)

os.makedirs("./temp", exist_ok=True)
if not minio_client.bucket_exists(bucket_name):
    minio_client.make_bucket(bucket_name)

def convert_wmf_to_png(wmf_path):
    """Convert WMF to PNG using Wand"""
    try:
        from wand.image import Image
        png_path = wmf_path.replace('.wmf', '.png')
        with Image(filename=wmf_path) as img:
            img.format = "png"
            img.save(filename=png_path)
        return png_path if os.path.exists(png_path) else None
    except Exception as e:
        print(f"WMF conversion failed: {e}")
        return None

def save_to_minio(local_path, object_name):
    """Upload to MinIO and return URL"""
    try:
        minio_client.fput_object(bucket_name, object_name, local_path)
        return f"{minio_external_url}/{bucket_name}/{object_name}"
    except Exception as e:
        print(f"MinIO upload failed: {e}")
        return None


def extract_images_with_positions(local_file_path: str) -> Dict[str, ImageInfo]:
    """提取图片并记录位置信息 - 跳过第一张图片"""
    images = {}
    
    try:
        import zipfile
        from xml.etree import ElementTree as ET
        
        with zipfile.ZipFile(local_file_path, 'r') as docx_zip:
            # 提取图片文件
            image_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]
            print(f"🖼️ 发现 {len(image_files)} 个图片文件")
            
            # 解析document.xml获取图片位置
            doc_xml = docx_zip.read('word/document.xml').decode('utf-8')
            
            # 解析关系文件获取图片引用
            try:
                rels_xml = docx_zip.read('word/_rels/document.xml.rels').decode('utf-8')
                rels_root = ET.fromstring(rels_xml)
                
                # 建立rId到文件的映射
                rid_to_file = {}
                for rel in rels_root.findall('.//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
                    rel_type = rel.get('Type', '')
                    if 'image' in rel_type:
                        rid_to_file[rel.get('Id')] = rel.get('Target')
                        
            except Exception as e:
                print(f"⚠️ 无法解析关系文件: {e}")
                rid_to_file = {}
            
            valid_image_count = 0
            
            # 处理每个图片文件
            for img_index, img_path in enumerate(image_files):
                try:
                    img_data = docx_zip.read(img_path)
                    
                    # 🚫 跳过第一张图片 (通常是Logo)
                    if img_index == 0:
                        print(f"🚫 跳过第一张图片: {img_path}")
                        continue
                    
                    img_extension = os.path.splitext(img_path)[1] or '.jpg'
                    img_filename = f"docx_image_{uuid.uuid4()}{img_extension}"
                    local_img_path = f"./temp/{img_filename}"
                    
                    with open(local_img_path, 'wb') as img_file:
                        img_file.write(img_data)
                    
                    # Convert WMF if needed
                    final_img_path = local_img_path
                    if img_extension.lower() == '.wmf':
                        png_path = convert_wmf_to_png(local_img_path)
                        if png_path:
                            final_img_path = png_path
                            img_filename = img_filename.replace('.wmf', '.png')
                    
                    # Upload to MinIO
                    minio_path = f"docx_images/{img_filename}"
                    image_url = save_to_minio(final_img_path, minio_path)
                    
                    if image_url:
                        valid_image_count += 1
                        image_id = f"img_{valid_image_count:03d}"  # 使用有效图片计数
                        
                        # 创建ImageInfo对象
                        image_info = ImageInfo(
                            id=image_id,
                            filename=img_filename,
                            url=image_url,
                            type="docx_embedded",
                            original_path=img_path,
                            paragraph_index=0,  # 将在后续步骤中更新
                            position_in_paragraph=0,
                            size={"width": 0, "height": 0}  # 将在下面更新
                        )
                        
                        # 尝试获取图片尺寸
                        try:
                            from PIL import Image
                            with Image.open(final_img_path) as img:
                                image_info.size = {"width": img.size[0], "height": img.size[1]}
                                print(f"✅ {image_id}: {img_filename} ({img.size[0]}x{img.size[1]})")
                        except:
                            print(f"✅ {image_id}: {img_filename}")
                        
                        images[image_id] = image_info
                    
                    # Cleanup
                    for path in [local_img_path, final_img_path]:
                        try:
                            if os.path.exists(path):
                                os.remove(path)
                        except:
                            pass
                            
                except Exception as e:
                    print(f"❌ Error processing {img_path}: {e}")
    
    except Exception as e:
        print(f"❌ Error extracting images: {e}")
    
    print(f"✅ 过滤后保留 {len(images)} 张有效图片")
    return images

def analyze_document_structure(doc: Document) -> DoclingDocument:
    """分析文档结构并创建DoclingDocument - 改进版本"""
    docling_doc = DoclingDocument()
    
    # 设置元数据
    docling_doc.metadata = {
        "extraction_time": datetime.now().isoformat(),
        "total_paragraphs": len(doc.paragraphs),
        "total_tables": len(doc.tables),
        "document_type": "docx"
    }
    
    print("🔍 开始文档结构分析...")
    
    # 分析段落并创建内容块
    sequence = 0
    current_question_block = None
    classification_stats = {}
    
    for para_index, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
            
        sequence += 1
        block_id = f"block_{sequence:03d}"
        
        # 使用改进的分类系统
        block_type = classify_block_type(text)
        
        # 统计分类结果
        classification_stats[block_type] = classification_stats.get(block_type, 0) + 1
        
        # 设置父块关系
        parent_block = None
        if block_type == "option" and current_question_block:
            parent_block = current_question_block
        elif block_type == "sub_question" and current_question_block:
            parent_block = current_question_block
        elif block_type == "answer" and current_question_block:
            parent_block = current_question_block
            
        # 更新当前题目块
        if block_type == "question":
            current_question_block = block_id
            print(f"📝 题目: {text[:60]}...")
        elif block_type == "option":
            print(f"  📋 选项: {text[:40]}...")
        elif block_type == "section_title":
            print(f"📂 标题: {text}")
            current_question_block = None  # 重置当前题目
        elif block_type == "answer":
            print(f"💡 答案: {text[:40]}...")
        elif block_type == "sub_question":
            print(f"  📝 子题: {text[:40]}...")
            
        # 创建内容块
        content_block = ContentBlock(
            id=block_id,
            type=block_type,
            sequence=sequence,
            text=text,
            paragraph_index=para_index,
            style=get_paragraph_style(para),
            parent_block=parent_block
        )
        
        docling_doc.add_content_block(content_block)
    
    # 输出分类统计
    print(f"\n📊 文档结构分析完成:")
    for block_type, count in sorted(classification_stats.items()):
        print(f"  {block_type}: {count} 个")
        
    return docling_doc

def is_real_question(text: str, prev_text: str = None) -> bool:
    """准确识别真实题目 - 排除干扰项"""
    if not text or not text.strip():
        return False
        
    text = text.strip()
    
    # 🚫 排除选项
    if re.match(r'^[A-D][．.、]\s*', text):
        return False
        
    # 🚫 排除题型标题
    if re.match(r'^[一二三四五六七八九十]+[、．]\s*(选择题|填空题|解答题|判断题|简答题|计算题)', text):
        return False
        
    # 🚫 排除章节标题
    if re.match(r'^第[一二三四五六七八九十\d]+[章节部分]\s*', text):
        return False
        
    # 🚫 排除答案解析
    if text.startswith(('解：', '证明：', '答：', '[解析]', '解析：', '分析：')):
        return False
        
    # 🚫 排除纯标题性文字
    title_keywords = ['考试', '测试', '练习', '作业', '复习', '总复习', '期中', '期末', '月考']
    if any(keyword in text and len(text) < 20 for keyword in title_keywords):
        return False
        
    # ✅ 识别真实题目模式
    question_patterns = [
        r'^\d+[．.、]\s*',           # 1. 或 1、
        r'^[(（]\d+[)）]\s*',        # (1) 或 （1）
        r'^第\s*\d+\s*题\s*',        # 第1题
        r'^题\s*\d+\s*[．.、]?\s*',   # 题1.
    ]
    
    for pattern in question_patterns:
        if re.match(pattern, text):
            # 进一步验证：确保后面有实际内容
            content_part = re.sub(pattern, '', text).strip()
            if len(content_part) > 5:  # 至少有5个字符的内容
                return True
    
    # ✅ 检查题目关键词（但需要有一定长度）
    if len(text) > 10:  # 避免标题类短文本
        question_keywords = ['下列', '以下', '关于', '设有', '已知', '如图', '求解', '计算', '证明', 
                           '选择', '判断', '填空', '简述', '分析', '解释', '说明']
        if any(keyword in text for keyword in question_keywords):
            return True
            
    return False

def is_section_title(text: str) -> bool:
    """判断是否为章节标题"""
    if not text or not text.strip():
        return False
        
    text = text.strip()
    
    # 题型标题
    if re.match(r'^[一二三四五六七八九十]+[、．]\s*(选择题|填空题|解答题|判断题|简答题|计算题)', text):
        return True
        
    # 章节标题  
    if re.match(r'^第[一二三四五六七八九十\d]+[章节部分]\s*', text):
        return True
        
    # 短标题
    if len(text) < 15 and any(keyword in text for keyword in ['考试', '测试', '练习', '作业', '复习']):
        return True
        
    return False

def is_answer_block(text: str) -> bool:
    """判断是否为答案解析块"""
    if not text or not text.strip():
        return False
        
    text = text.strip()
    
    answer_prefixes = ['解：', '证明：', '答：', '[解析]', '解析：', '分析：', '参考答案：']
    return any(text.startswith(prefix) for prefix in answer_prefixes)

def is_sub_question(text: str) -> bool:
    """判断是否为子问题"""
    if not text or not text.strip():
        return False
        
    text = text.strip()
    
    # (1) (2) 形式，但不是主题目开头
    if re.match(r'^[(（][1-9][)）]', text) and not re.match(r'^[(（]1[)）]', text):
        return True
        
    # 子问题关键词
    sub_patterns = [
        r'^请[问回答计算解释说明]',
        r'^试[问求计算证明分析]',  
        r'^若[果]',
        r'^当.*时'
    ]
    
    return any(re.match(pattern, text) for pattern in sub_patterns)

def classify_block_type(text: str, context: Dict = None) -> str:
    """精确分类每个文本块的类型 - 按优先级排序"""
    if not text or not text.strip():
        return 'empty'
        
    text = text.strip()
    
    # 优先级从高到低
    if is_section_title(text):
        return 'section_title'  # 不分配图片
        
    if is_option_block(text):
        return 'option'  # 可能有图片，但要关联到父题目
        
    if is_answer_block(text):
        return 'answer'  # 可能有解析图
        
    if is_sub_question(text):
        return 'sub_question'  # 继承父题目的图片
        
    if is_real_question(text):
        return 'question'  # 可能有图片
        
    return 'text'  # 普通段落，通常不分配图片

def is_option_block(text: str) -> bool:
    """判断是否为选项块"""
    option_patterns = [
        r'^\s*[A-D][\.、]\s*',     # A. B. C. D.
        r'^\s*\([A-D]\)\s*',       # (A) (B) (C) (D)
        r'^\s*[①②③④⑤⑥⑦⑧]\s*',   # ①②③④
    ]
    
    for pattern in option_patterns:
        if re.match(pattern, text):
            return True
            
    return False

def get_paragraph_style(para) -> Optional[str]:
    """获取段落样式"""
    try:
        if para.style:
            return para.style.name
    except:
        pass
    return None



def smart_image_assignment(docling_doc: DoclingDocument, images: Dict[str, ImageInfo]):
    """通用图片分配算法 - 支持各种图片引用模式"""
    
    print("🎯 开始通用图片分配算法...")
    
    # Step 1: 定义常见的图片引用模式
    image_reference_patterns = [
        r'图\s*(\d+)',                    # 图1, 图2, 图 3
        r'图\s*(\d+)[-－](\d+)',         # 图5-1, 图5-2
        r'如图\s*(\d+)',                 # 如图1, 如图2
        r'见图\s*(\d+)',                 # 见图1, 见图2  
        r'参见图\s*(\d+)',               # 参见图1, 参见图2
        r'附图\s*([A-Za-z]+)',          # 附图A, 附图B
        r'示意图\s*(\d+)',               # 示意图1, 示意图2
        r'Figure\s*(\d+)',              # Figure 1, Figure 2
        r'Fig\.\s*(\d+)',               # Fig.1, Fig.2
        r'[图图]\s*([K-Z])[-－]?(\d+)[-－](\d+)',  # 图K-19-1, 图A-5-2
        r'第\s*(\d+)\s*题图',            # 第1题图, 第2题图
    ]
    
    # Step 2: 找到所有有图片引用的题目
    questions_with_images = []  # [(block, matched_patterns, sequence)]
    question_blocks = [block for block in docling_doc.content_blocks if block.type == 'question']
    
    print(f"\n📊 基本信息:")
    print(f"  📝 题目数量: {len(question_blocks)}")
    print(f"  🖼️ 图片数量: {len(images)}")
    
    print(f"\n🔍 分析图片引用模式:")
    for block in question_blocks:
        matched_patterns = []
        
        # 尝试所有图片引用模式
        for pattern in image_reference_patterns:
            matches = re.findall(pattern, block.text)
            if matches:
                matched_patterns.extend(matches)
        
        if matched_patterns:
            questions_with_images.append((block, matched_patterns, block.sequence))
            # 显示找到的所有引用
            refs_str = ', '.join([str(m) if isinstance(m, (str, int)) else '-'.join(map(str, m)) for m in matched_patterns])
            print(f"  📍 {block.id}(seq:{block.sequence}): 发现引用 [{refs_str}]")
        else:
            print(f"  ⚪ {block.id}: 无图片引用")
    
    # Step 3: 按文档位置排序 (sequence)
    questions_with_images.sort(key=lambda x: x[2])
    image_items = list(images.items())
    
    # Step 4: 智能匹配策略
    print(f"\n🎯 智能分配策略:")
    question_image_mapping = {}
    used_images = set()
    
    # 检查是否有有效图片
    if len(image_items) == 0:
        print("  ⚠️ 没有有效图片可分配 (可能都被过滤掉了)")
        return question_image_mapping
    
    if len(questions_with_images) == 0:
        print("  ⚠️ 没有需要图片的题目")
        return question_image_mapping
    
    # 策略1: 尝试精确数字匹配 (如果可能)
    if len(questions_with_images) <= len(images):
        print("  📋 策略: 顺序分配 (题目≤图片)")
        for i, (block, patterns, seq) in enumerate(questions_with_images):
            if i < len(image_items):
                image_id = image_items[i][0]
                docling_doc.link_image_to_block(block.id, image_id)
                question_image_mapping[block.id] = image_id
                used_images.add(image_id)
                print(f"    ✅ {block.id} → {image_id} (第{i+1}对)")
    else:
        print("  📋 策略: 比例分配 (题目>图片)")
        # 如果题目多于图片，尽可能均匀分配
        for i, (block, patterns, seq) in enumerate(questions_with_images):
            if len(image_items) > 0:  # 确保有图片可分配
                image_index = i * len(image_items) // len(questions_with_images)
                image_id = image_items[image_index][0]
                
                if image_id not in used_images:
                    docling_doc.link_image_to_block(block.id, image_id)
                    question_image_mapping[block.id] = image_id
                    used_images.add(image_id)
                    print(f"    ✅ {block.id} → {image_id} (比例分配)")
    
    # Step 5: 统计结果
    unused_images = [img_id for img_id, _ in image_items if img_id not in used_images]
    
    print(f"\n📊 分配统计:")
    print(f"  🎯 有图片引用的题目: {len(questions_with_images)} 个")
    print(f"  ✅ 成功分配: {len(question_image_mapping)} 个")
    print(f"  🖼️ 已使用图片: {len(used_images)}/{len(images)}")
    print(f"  ⚠️ 未使用图片: {len(unused_images)} 张")
    
    if unused_images:
        print(f"  📋 未使用图片: {unused_images}")
    
    return question_image_mapping


def link_images_to_blocks(docling_doc: DoclingDocument, images: Dict[str, ImageInfo]):
    """建立图片与内容块的关联关系"""
    print("🔍 使用智能匹配策略...")
    question_mapping = smart_image_assignment(docling_doc, images)
    return question_mapping

@app.post("/api/extract_questions_with_images")
async def extract_questions(
    file: UploadFile = File(...), 
    legacy_format: bool = True
):
    try:
        file_ext = os.path.splitext(file.filename)[1].lower()
        local_file_path = f"./temp/{uuid.uuid4()}{file_ext}"
        
        with open(local_file_path, "wb") as f:
            shutil.copyfileobj(file.file, f)

        if file_ext != ".docx":
            return {"error": "Only DOCX files supported in this version"}

        print("📂 开始结构化文档提取...")
        
        # STEP 1: 提取图片并记录位置信息
        print("🖼️ 提取图片...")
        images = extract_images_with_positions(local_file_path)
        
        # STEP 2: 分析文档结构
        print("📄 分析文档结构...")
        doc = Document(local_file_path)
        docling_doc = analyze_document_structure(doc)
        
        # STEP 3: 建立图片与内容块的关联
        print("🔗 建立图片关联...")
        link_images_to_blocks(docling_doc, images)
        
        # STEP 4: 将图片信息加入DoclingDocument
        for image_id, image_info in images.items():
            docling_doc.add_image(image_info)
        
        print(f"✅ 结构化提取完成:")
        print(f"   📝 内容块数量: {len(docling_doc.content_blocks)}")
        print(f"   🖼️ 图片数量: {len(docling_doc.images)}")
        print(f"   🔗 关联关系: {len(docling_doc.relationships)}")
        
        # Cleanup
        try:
            os.remove(local_file_path)
        except:
            pass
        
        # 返回格式选择
        if legacy_format:
            # 向后兼容：返回原格式
            print("📋 返回兼容格式...")
            return convert_to_legacy_format(docling_doc)
        else:
            # 返回新的结构化格式
            print("📋 返回结构化格式...")
            return docling_doc.to_dict()
        
    except Exception as e:
        return {"error": f"Processing error: {str(e)}"}

def convert_to_legacy_format(docling_doc: DoclingDocument) -> Dict[str, Any]:
    """转换DoclingDocument为兼容的旧格式"""
    
    print("🔄 开始转换为兼容格式...")
    
    # 重建文本内容
    text_parts = []
    total_markers_added = 0
    
    for block in sorted(docling_doc.content_blocks, key=lambda x: x.sequence):
        block_text = block.text
        original_block_text = block_text
        
        # 在文本中插入图片标记（插入到K-19-X引用位置）
        if block.image_refs:
            print(f"📝 处理块 {block.id} (类型: {block.type}):")
            print(f"  原文本: {original_block_text[:100]}...")
            print(f"  图片引用: {block.image_refs}")
            
            for image_ref in block.image_refs:
                if image_ref in docling_doc.images:
                    # 使用新格式的图片标记
                    marker = f"{{{{{image_ref.upper()}}}}}"
                    
                    # 在图片引用位置插入标记，支持多种引用模式
                    image_ref_patterns = [
                        r'(图\s*\d+)',                     # 图1, 图2
                        r'(图\s*\d+[-－]\d+)',            # 图5-1, 图5-2  
                        r'(如图\s*\d+)',                  # 如图1, 如图2
                        r'(见图\s*\d+)',                  # 见图1, 见图2
                        r'(参见图\s*\d+)',                # 参见图1, 参见图2
                        r'(附图\s*[A-Za-z]+)',           # 附图A, 附图B
                        r'(示意图\s*\d+)',                # 示意图1, 示意图2
                        r'(Figure\s*\d+)',               # Figure 1, Figure 2
                        r'(Fig\.\s*\d+)',                # Fig.1, Fig.2
                        r'([图图]\s*[K-Z][-－]?\d+[-－]?\d+)',  # 图K-19-1, 图A-5-2
                        r'(第\s*\d+\s*题图)',             # 第1题图, 第2题图
                    ]
                    
                    inserted = False
                    for pattern in image_ref_patterns:
                        if re.search(pattern, block_text):
                            # 在第一个图片引用后面插入标记
                            block_text = re.sub(pattern, f'\\1 {marker}', block_text, count=1)
                            print(f"  ✅ 在图片引用后插入标记: {marker}")
                            inserted = True
                            break
                    
                    if not inserted:
                        # 如果没有找到图片引用，插入到题目编号后
                        number_pattern = r'^(\d+\．)'
                        if re.search(number_pattern, block_text):
                            block_text = re.sub(number_pattern, f'\\1{marker} ', block_text, count=1)
                            print(f"  ✅ 在题目编号后插入标记: {marker}")
                        else:
                            # 最后备选：添加到末尾
                            block_text += f" {marker}"
                            print(f"  ✅ 在末尾添加标记: {marker}")
                    
                    total_markers_added += 1
                else:
                    print(f"  ❌ 图片引用不存在: {image_ref}")
        
        # 更新块的文本以包含标记
        block.text = block_text
        text_parts.append(block_text)
    
    all_text = "\n".join(text_parts)
    
    # 🔍 修复验证：确保所有需要图的题目都有标记
    print(f"\n🔍 验证题目标记完整性:")
    question_blocks = [b for b in docling_doc.content_blocks if b.type == 'question']
    for i, block in enumerate(question_blocks, 1):
        # 检查题目是否有任何图片引用（这才是需要图片的标准）
        has_image_reference = bool(
            re.search(r'图\s*\d+', block.text) or
            re.search(r'图\s*\d+[-－]\d+', block.text) or  
            re.search(r'如图\s*\d+', block.text) or
            re.search(r'见图\s*\d+', block.text) or
            re.search(r'参见图\s*\d+', block.text) or
            re.search(r'附图\s*[A-Za-z]+', block.text) or
            re.search(r'示意图\s*\d+', block.text) or
            re.search(r'Figure\s*\d+', block.text) or
            re.search(r'Fig\.\s*\d+', block.text) or
            re.search(r'[图图]\s*[K-Z][-－]?\d+[-－]?\d+', block.text) or
            re.search(r'第\s*\d+\s*题图', block.text)
        )
        has_image_refs = bool(block.image_refs)
        has_marker_in_text = bool(re.search(r'\{\{[A-Z_0-9]+\}\}', block.text))  # 检查块文本中的标记
        
        # 正确的验证逻辑：有图片引用的题目应该有图片分配和标记
        status = "✅" if (has_image_reference == has_image_refs) else "❌"
        print(f"  题目{i} ({block.id}): 图片引用={has_image_reference}, 图片分配={has_image_refs} {status}")
        
        if has_image_reference and not has_image_refs:
            print(f"    ⚠️ 有图片引用但未分配图片的题目: {block.text[:50]}...")
        elif has_image_refs and not has_image_reference:
            print(f"    ⚠️ 无图片引用但分配了图片的题目: {block.text[:50]}...")
    
    # 🔍 DIAGNOSTIC: 检查最终文本中的标记
    print(f"\n🔍 兼容格式转换诊断:")
    print(f"  📝 内容块数量: {len(docling_doc.content_blocks)}")
    print(f"  🖼️ 图片数量: {len(docling_doc.images)}")
    print(f"  🏷️ 添加的标记数量: {total_markers_added}")
    
    # 检查文本中的标记
    marker_pattern = r'\{\{[A-Z_0-9]+\}\}'
    final_markers = re.findall(marker_pattern, all_text)
    print(f"  ✅ 文本中发现的标记: {len(final_markers)} 个")
    
    if final_markers:
        print(f"  📋 标记列表: {final_markers}")
        
        # 显示标记在文本中的位置
        print("\n🔍 标记位置分析:")
        for match in re.finditer(marker_pattern, all_text):
            start = max(0, match.start()-30)
            end = min(len(all_text), match.end()+30)
            context = all_text[start:end].replace('\n', '\\n')
            print(f"  {match.group()}: ...{context}...")
    else:
        print("⚠️ 警告: 最终文本中未发现任何图片标记!")
        print("📄 文本前500字符:")
        print(all_text[:500])
        
        # 额外诊断：检查是否有图片但没有生成标记
        if len(docling_doc.images) > 0:
            print(f"\n🔍 深度诊断 - 有图片({len(docling_doc.images)}张)但无标记:")
            for block in docling_doc.content_blocks:
                if block.image_refs:
                    print(f"  块 {block.id}: 有图片引用 {block.image_refs} 但未生成标记")
                    print(f"    文本: {block.text[:100]}...")
    
    # 📝 写入诊断日志
    debug_log_path = "./temp/legacy_conversion_debug.log"
    try:
        os.makedirs("./temp", exist_ok=True)
        with open(debug_log_path, "w", encoding="utf-8") as log_file:
            log_file.write("=== 兼容格式转换诊断 ===\n")
            log_file.write(f"转换时间: {datetime.now().isoformat()}\n")
            log_file.write(f"内容块数: {len(docling_doc.content_blocks)}\n")
            log_file.write(f"图片数: {len(docling_doc.images)}\n")
            log_file.write(f"添加标记数: {total_markers_added}\n")
            log_file.write(f"最终标记数: {len(final_markers)}\n\n")
            
            log_file.write("=== 内容块详情 ===\n")
            for block in sorted(docling_doc.content_blocks, key=lambda x: x.sequence):
                log_file.write(f"\n块ID: {block.id}\n")
                log_file.write(f"类型: {block.type}\n")
                log_file.write(f"图片引用: {block.image_refs}\n")
                log_file.write(f"文本: {block.text[:200]}...\n")
                log_file.write("-" * 50 + "\n")
            
            log_file.write("\n=== 最终文本（前2000字符） ===\n")
            log_file.write(all_text[:2000])
            
        print(f"📝 兼容格式诊断日志: {debug_log_path}")
    except Exception as e:
        print(f"⚠️ 写入诊断日志失败: {e}")
    
    # 转换图片格式
    legacy_images = []
    for image_id, image_info in docling_doc.images.items():
        legacy_images.append({
            "image_url": image_info.url,
            "image_type": image_info.type,
            "image_id": image_id.upper(),  # 保持大写格式一致性
            "original_path": image_info.original_path
        })
    
    return {
        "textContent": all_text,
        "images": legacy_images,
        "imageCount": len(legacy_images),
        "message": f"Extracted {len(legacy_images)} images with structured parsing",
        # 附加结构化信息（供调试使用）
        "structure_info": {
            "total_blocks": len(docling_doc.content_blocks),
            "question_blocks": len([b for b in docling_doc.content_blocks if b.type == "question"]),
            "option_blocks": len([b for b in docling_doc.content_blocks if b.type == "option"]),
            "relationships": len(docling_doc.relationships)
        }
    }

async def extract_questions_with_images_core(file: UploadFile = File(...)):
    """Core function for extracting questions with images - used by route_api.py"""
    return await extract_questions(file)

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8003)